"""Build the final AIDEOM-VN course report as a Word document.

The report is generated from local source-code outputs only:
CSV tables, PNG charts/screenshots, and the project files already present in
the repository. It intentionally avoids hard-coded empirical results that are
not present in the exported artifacts.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = ROOT / "reports" / "figures"
SCREEN_DIR = FIG_DIR / "screenshots"
OUT_PATH = ROOT / "reports" / "aideom_vn_bao_cao_cuoi_ky.docx"


def rel(path: Path) -> str:
    """Return a POSIX relative path for captions and source notes."""
    return path.relative_to(ROOT).as_posix()


def read_csv(name: str) -> pd.DataFrame:
    """Read one exported report artifact from reports/figures."""
    return pd.read_csv(FIG_DIR / name)


def fmt_num(value, digits: int = 2) -> str:
    """Format numbers compactly while preserving text values."""
    if pd.isna(value):
        return ""
    if isinstance(value, (int, float)):
        if abs(value) >= 1000:
            return f"{value:,.{digits}f}"
        return f"{value:.{digits}f}"
    return str(value)


def set_cell_shading(cell, fill: str) -> None:
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    """Set table cell text using the document font."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 16 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold: bool = False):
    """Add a paragraph with Times New Roman direct formatting."""
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    """Add a short bullet list using Word's list style."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    """Add a short numbered list using Word's list style."""
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_table_df(
    doc: Document,
    df: pd.DataFrame,
    caption: str,
    source: str,
    max_rows: int | None = None,
    digits: int = 2,
) -> None:
    """Insert a dataframe as a formatted Word table with caption and source."""
    shown = df.copy()
    if max_rows is not None:
        shown = shown.head(max_rows)
    add_paragraph(doc, caption, bold=True)
    table = doc.add_table(rows=1, cols=len(shown.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, col in enumerate(shown.columns):
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
        set_cell_text(table.rows[0].cells[i], str(col), bold=True)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(shown.columns):
            set_cell_text(cells[i], fmt_num(row[col], digits=digits))
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    run = note.add_run(f"Nguồn: {source}")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(89, 89, 89)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.3) -> None:
    """Insert a figure if it exists and add a source caption."""
    if not image_path.exists():
        add_paragraph(doc, f"[Thiếu hình: {rel(image_path)}]")
        return
    doc.add_picture(str(image_path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"{caption}\nNguồn hình: {rel(image_path)}")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(89, 89, 89)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading with Times New Roman formatting."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.color.rgb = RGBColor(31, 78, 121)
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)


def add_callout(doc: Document, title: str, body: str) -> None:
    """Add a restrained academic callout as a one-row table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF5FF")
    cell.text = ""
    p = cell.paragraphs[0]
    r1 = p.add_run(title + ": ")
    r1.bold = True
    r2 = p.add_run(body)
    for run in (r1, r2):
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(11)


def configure_document(doc: Document) -> None:
    """Set global Word styles and page layout."""
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    # Front matter intentionally has no page number. Pagination starts at
    # Chapter 1 in a separate section created by start_numbered_section().
    footer = section.footer.paragraphs[0]
    footer.clear()


def start_numbered_section(doc: Document, start: int = 1):
    """Start a new-page section with centered PAGE numbering."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.footer.is_linked_to_previous = False

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    field_text = OxmlElement("w:t")
    field_text.text = str(start)
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.extend([field_begin, instruction, field_separate, field_text, field_end])

    page_number_type = OxmlElement("w:pgNumType")
    page_number_type.set(qn("w:start"), str(start))
    section._sectPr.append(page_number_type)
    return section


def summarize_inputs() -> pd.DataFrame:
    """Create a compact input-data overview table from actual CSV files."""
    rows = []
    for name in [
        "vietnam_macro_2020_2025.csv",
        "vietnam_sectors_2024.csv",
        "vietnam_regions_2024.csv",
    ]:
        df = pd.read_csv(ROOT / "data" / name)
        rows.append(
            {
                "Tệp dữ liệu": name,
                "Số dòng": len(df),
                "Số cột": len(df.columns),
                "Một số biến chính": ", ".join(df.columns[:6]),
            }
        )
    rows.append(
        {
            "Tệp dữ liệu": "Tham số chính sách trong src/",
            "Số dòng": "n/a",
            "Số cột": "n/a",
            "Một số biến chính": "SCENARIOS, hệ số beta, ràng buộc LP/MIP, action mapping RL",
        }
    )
    return pd.DataFrame(rows)


def module_table() -> pd.DataFrame:
    """Create a concise 12-module overview table from source/output names."""
    return pd.DataFrame(
        [
            ["Bài 1", "Cobb-Douglas AI", "Ước lượng TFP và dự báo GDP 2030", "src/bai01_cobb_douglas.py"],
            ["Bài 2", "LP ngân sách số", "Tối ưu bốn hạng mục K/D/AI/H", "src/bai02_lp_budget.py"],
            ["Bài 3", "Priority ngành", "Xếp hạng 10 ngành bằng chuẩn hóa min-max", "src/bai03_priority.py"],
            ["Bài 4", "LP ngành-vùng", "Phân bổ 6 vùng x 4 hạng mục có công bằng", "src/bai04_region_lp.py"],
            ["Bài 5", "MIP dự án", "Chọn danh mục 15 dự án chuyển đổi số", "src/bai05_mip_projects.py"],
            ["Bài 6", "TOPSIS vùng", "Xếp hạng 6 vùng đầu tư AI", "src/bai06_topsis.py"],
            ["Bài 7", "Pareto NSGA-II/fallback", "Tối ưu đa mục tiêu tăng trưởng-rủi ro", "src/bai07_pareto.py"],
            ["Bài 8", "Tối ưu động", "Mô phỏng 2026-2035 bằng SLSQP", "src/bai08_dynamic.py"],
            ["Bài 9", "Lao động AI", "Tối ưu AI/đào tạo theo ngành", "src/bai09_labor_ai.py"],
            ["Bài 10", "Stochastic SP", "Quy hoạch ngẫu nhiên hai giai đoạn", "src/bai10_stochastic.py"],
            ["Bài 11", "Q-learning", "Chính sách thích nghi 81 trạng thái", "src/bai11_q_learning.py"],
            ["Bài 12", "AIDEOM tích hợp", "So sánh 5 kịch bản chính sách", "src/scenario_engine.py"],
        ],
        columns=["Mã", "Tên module", "Vai trò", "Source code"],
    )


def add_cover(doc: Document) -> None:
    """Create the first page of the report."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRƯỜNG/ĐƠN VỊ ĐÀO TẠO\nMÔN: CÁC MÔ HÌNH RA QUYẾT ĐỊNH")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(13)
    r.bold = True

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO CUỐI KỲ")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AIDEOM-VN: Dashboard mô hình ra quyết định cho kinh tế số, AI và phân bổ nguồn lực tại Việt Nam")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(15)
    r.bold = True

    doc.add_paragraph("\n")
    add_callout(
        doc,
        "Phạm vi báo cáo",
        "Báo cáo sử dụng kết quả đã xuất trong reports/figures, dữ liệu CSV trong data/ và logic mô hình trong src/. "
        "Các số liệu định lượng được trình bày đều có nguồn từ các artifact cục bộ của dự án.",
    )
    doc.add_paragraph("\n\n")
    for line in [
        "Nhóm thực hiện: AIDEOM-VN",
        "Nền tảng triển khai: Python, Streamlit, pandas/numpy, Plotly, PuLP/SciPy và các fallback mô hình",
        "Ngày tạo báo cáo: 31/05/2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    doc.add_page_break()


def build_report() -> None:
    """Build the final report document."""
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    macro = read_csv("00_input_macro.csv")
    bai01_fit = read_csv("01_bai01_cobb_douglas_fit.csv")
    bai01_forecast = read_csv("01_bai01_forecast_2030.csv")
    bai02_alloc = read_csv("02_bai02_allocation.csv")
    bai03_rank = read_csv("03_bai03_priority_ranking.csv")
    bai04_matrix = read_csv("04_bai04_allocation_matrix.csv")
    bai04_fair = read_csv("04_bai04_fairness_comparison.csv")
    bai05_selected = read_csv("05_bai05_selected_projects.csv")
    bai06_topsis = read_csv("06_bai06_topsis_expert.csv")
    bai07_compare = read_csv("07_bai07_growth_vs_compromise.csv")
    bai08_summary = read_csv("08_bai08_strategy_summary.csv")
    bai09_jobs = read_csv("09_bai09_allocation_jobs.csv")
    bai10_vss = read_csv("10_bai10_vss_evpi.csv")
    bai11_compare = read_csv("11_bai11_policy_comparison.csv")
    bai12_kpi = read_csv("12_bai12_aideom_kpi.csv")
    bai12_alloc = read_csv("12_bai12_allocation_long.csv")
    bai12_recs = read_csv("12_bai12_policy_recommendations.csv")
    run_summary = read_csv("00_model_run_summary.csv")

    mape = float(run_summary.loc[run_summary["module"] == "Bài 1", "main_result"].iloc[0].split("=")[1])
    best_scenario = bai12_kpi.sort_values("Overall_score", ascending=False).iloc[0]
    highest_gdp = bai12_kpi.sort_values("GDP_gain", ascending=False).iloc[0]
    highest_netjob = bai12_kpi.sort_values("NetJob", ascending=False).iloc[0]

    add_heading(doc, "MỤC LỤC TÓM TẮT", 1)
    add_paragraph(
        doc,
        "Báo cáo gồm bảy chương chính và phần phụ lục. Do tài liệu được tạo tự động từ "
        "kết quả chạy phần mềm, người đọc có thể đối chiếu từng bảng/hình với tệp CSV hoặc PNG "
        "trong thư mục reports/figures của dự án.",
    )
    add_numbered(
        doc,
        [
            "Chương 1 giới thiệu bối cảnh, mục tiêu, câu hỏi và phạm vi nghiên cứu.",
            "Chương 2 trình bày cơ sở lý thuyết và các mô hình ra quyết định.",
            "Chương 3 mô tả dữ liệu, tiền xử lý, pipeline và thiết kế webapp.",
            "Chương 4 tổng hợp kết quả chạy các module và giao diện dashboard.",
            "Chương 5 so sánh năm kịch bản chính sách của AIDEOM-VN.",
            "Chương 6 thảo luận ý nghĩa chính sách, hạn chế và cải tiến.",
            "Chương 7 kết luận, khuyến nghị và hướng phát triển tiếp theo.",
        ],
    )
    start_numbered_section(doc, start=1)

    add_heading(doc, "CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Bối cảnh nghiên cứu", 2)
    add_paragraph(
        doc,
        "Việt Nam đang chuyển từ mô hình tăng trưởng dựa nhiều vào vốn vật chất, lao động và "
        "mở rộng sản xuất sang mô hình tăng trưởng có hàm lượng công nghệ, dữ liệu và năng suất cao hơn. "
        "Trong bối cảnh kinh tế số, dữ liệu, trí tuệ nhân tạo và năng lực nhân lực số trở thành yếu tố đầu vào "
        "quan trọng, bài toán chính sách không chỉ là tăng quy mô đầu tư mà còn là phân bổ nguồn lực đúng nơi, "
        "đúng thời điểm và kiểm soát các rủi ro phát sinh.",
    )
    add_paragraph(
        doc,
        "Dữ liệu vĩ mô trong dự án cho thấy giai đoạn 2020-2025 có GDP theo giá trị nghìn tỷ VND tăng từ "
        f"{macro['GDP_trillion_VND'].iloc[0]:,.1f} lên {macro['GDP_trillion_VND'].iloc[-1]:,.1f}; "
        f"tỷ trọng kinh tế số trong GDP tăng từ {macro['digital_economy_share_GDP_pct'].iloc[0]:.1f}% "
        f"lên {macro['digital_economy_share_GDP_pct'].iloc[-1]:.1f}%. Các con số này được lấy trực tiếp "
        "từ vietnam_macro_2020_2025.csv và là cơ sở để nhóm xây dựng nguyên mẫu dashboard AIDEOM-VN.",
    )
    add_heading(doc, "1.2. Mục tiêu nghiên cứu", 2)
    add_paragraph(
        doc,
        "Mục tiêu của nhóm là xây dựng một nguyên mẫu AIDEOM-VN có thể tích hợp dữ liệu, mô hình định lượng "
        "và trực quan hóa tương tác nhằm hỗ trợ phân tích chính sách kinh tế số và AI tại Việt Nam.",
    )
    add_bullets(
        doc,
        [
            "Dự báo kinh tế vĩ mô và mô phỏng GDP đến năm 2030 bằng hàm sản xuất mở rộng.",
            "Đánh giá mức độ sẵn sàng số theo ngành và vùng dựa trên nhiều tiêu chí.",
            "Tối ưu phân bổ nguồn lực bằng LP, MIP, quy hoạch động, quy hoạch ngẫu nhiên và Pareto.",
            "Mô phỏng tác động của AI tới việc làm mới, việc làm nâng cấp, thay thế và việc làm ròng.",
            "Đánh giá các rủi ro bất bình đẳng vùng, an ninh dữ liệu, phát thải và lao động.",
            "Trực quan hóa kết quả bằng webapp Streamlit để người dùng thao tác, tải bảng và xem khuyến nghị.",
        ],
    )
    add_heading(doc, "1.3. Câu hỏi nghiên cứu", 2)
    add_bullets(
        doc,
        [
            "Kịch bản chính sách nào tạo ra kết quả tăng trưởng/GDP gain cao nhất trong mô hình tích hợp?",
            "Kịch bản nào cân bằng tốt nhất giữa tăng trưởng, lao động, rủi ro và công bằng vùng?",
            "Dashboard hỗ trợ người dùng ra quyết định như thế nào thông qua bảng, biểu đồ và cảnh báo?",
        ],
    )
    add_heading(doc, "1.4. Phạm vi nghiên cứu", 2)
    add_paragraph(
        doc,
        "Phạm vi dữ liệu định lượng của dự án bao gồm dữ liệu vĩ mô Việt Nam giai đoạn 2020-2025, dữ liệu "
        "10 ngành kinh tế năm 2024 và dữ liệu 6 vùng kinh tế năm 2024. Các kịch bản chính sách trong Bài 12 "
        "sử dụng ngân sách mặc định 50.000 theo đơn vị mô hình, chia cho bốn nhóm K, D, AI và H. Phần dự báo "
        "Cobb-Douglas tạo quỹ đạo đến năm 2030; phần tối ưu động mô phỏng 2026-2035.",
    )
    add_heading(doc, "1.5. Cấu trúc báo cáo", 2)
    add_paragraph(
        doc,
        "Các chương tiếp theo lần lượt trình bày nền tảng lý thuyết, dữ liệu và thiết kế hệ thống, kết quả chạy "
        "từng module, so sánh năm kịch bản, thảo luận chính sách, kết luận và phụ lục kỹ thuật.",
    )
    doc.add_page_break()

    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ MÔ HÌNH RA QUYẾT ĐỊNH", 1)
    add_heading(doc, "2.1. Khung mô hình AIDEOM-VN", 2)
    add_paragraph(
        doc,
        "AIDEOM-VN được thiết kế như một hệ thống gồm sáu nhóm chức năng. Sáu nhóm này không thay thế "
        "mười hai bài toán chi tiết mà gom chúng thành các lớp phân tích dễ hiểu cho người dùng kinh tế.",
    )
    add_table_df(
        doc,
        pd.DataFrame(
            [
                ["M1", "Dự báo kinh tế", "Bài 1, Bài 8", "GDP, TFP, quỹ đạo 2026-2035"],
                ["M2", "Đánh giá sẵn sàng số", "Bài 3, Bài 6", "Priority ngành, TOPSIS vùng"],
                ["M3", "Tối ưu phân bổ", "Bài 2, Bài 4, Bài 5, Bài 7, Bài 10", "Phân bổ ngân sách, dự án, Pareto, SP"],
                ["M4", "Mô phỏng lao động", "Bài 9", "NewJob, UpgradeJob, DisplacedJob, NetJob"],
                ["M5", "Đánh giá rủi ro", "Bài 7, Bài 10, Bài 12", "Cyber, emission, inequality, uncertainty"],
                ["M6", "Dashboard ra quyết định", "app.py và pages/", "Streamlit, tabs, KPI cards, download CSV"],
            ],
            columns=["Module", "Nhóm chức năng", "Bài liên quan", "Kết quả chính"],
        ),
        "Bảng 1. Khung sáu module của hệ thống AIDEOM-VN",
        "Tổng hợp từ cấu trúc source code trong src/ và pages/",
    )
    add_figure(doc, FIG_DIR / "12_bai12_overall_score.png", "Hình 1. Sơ đồ/biểu đồ tổng hợp Overall_score trong module tích hợp", 5.8)

    add_heading(doc, "2.2. Hàm sản xuất Cobb-Douglas mở rộng", 2)
    add_paragraph(
        doc,
        "Bài 1 sử dụng hàm sản xuất Cobb-Douglas mở rộng: Y_t = A_t K_t^alpha L_t^beta D_t^gamma "
        "AI_t^delta H_t^theta. Trong đó Y là GDP, K là vốn, L là lao động hoặc năng suất/lao động theo dữ liệu "
        "linh hoạt, D là tỷ trọng kinh tế số, AI là mức ứng dụng AI, H là vốn nhân lực, còn A_t là năng suất nhân tố "
        "tổng hợp được suy ra từ dữ liệu. Bộ hệ số mặc định trong source code là alpha=0,33; beta=0,42; gamma=0,10; "
        "delta=0,08; theta=0,07.",
    )
    add_heading(doc, "2.3. Mô hình tối ưu hóa phân bổ nguồn lực", 2)
    add_paragraph(
        doc,
        "Nhóm mô hình tối ưu hóa bao gồm LP cho ngân sách số, LP theo vùng-hạng mục, MIP lựa chọn dự án, "
        "quy hoạch ngẫu nhiên hai giai đoạn và tối ưu đa mục tiêu. Điểm chung là các mô hình đều xác định biến "
        "quyết định, hàm mục tiêu và ràng buộc nguồn lực. Ví dụ, Bài 2 tối đa hóa Z = 0,85x1 + 1,20x2 + 0,95x3 + "
        "1,35x4 với ràng buộc tổng ngân sách và các mức đầu tư tối thiểu.",
    )
    add_heading(doc, "2.4. Phương pháp TOPSIS đánh giá mức độ ưu tiên", 2)
    add_paragraph(
        doc,
        "TOPSIS xếp hạng đối tượng theo khoảng cách tới nghiệm lý tưởng tốt nhất và nghiệm lý tưởng xấu nhất. "
        "Trong Bài 6, bảy tiêu chí đầu như GRDP bình quân, FDI, chỉ số số, AI readiness, lao động đào tạo, R&D và "
        "internet penetration được xem là tiêu chí lợi ích; gini_coef là tiêu chí chi phí. Source code cũng có "
        "trọng số entropy để so sánh với trọng số chuyên gia.",
    )
    add_heading(doc, "2.5. Mô hình mô phỏng tác động lao động", 2)
    add_paragraph(
        doc,
        "Bài 9 mô hình hóa tác động của AI tới thị trường lao động bằng bốn đại lượng: việc làm mới do đầu tư AI, "
        "việc làm được nâng cấp do đầu tư đào tạo, việc làm bị thay thế do rủi ro tự động hóa và năng lực đào tạo lại. "
        "Việc làm ròng được tính theo NetJob_i = NewJob_i + UpgradeJob_i - DisplacedJob_i, kèm ràng buộc NetJob_i không âm "
        "và DisplacedJob_i không vượt RetrainingCapacity_i.",
    )
    add_heading(doc, "2.6. Mô hình đánh giá rủi ro và kịch bản", 2)
    add_paragraph(
        doc,
        "Bài 12 tổng hợp rủi ro thành ba KPI chính gồm Inequality_risk, Cyber_risk và Emission_risk. Các rủi ro này "
        "được mô phỏng từ tỷ trọng phân bổ K, D, AI và H trong từng kịch bản. Bài 7 bổ sung cách nhìn Pareto qua bốn mục "
        "tiêu GDP gain, inequality, emission và net cyber risk; Bài 10 mô phỏng bất định bằng bốn kịch bản xác suất.",
    )
    doc.add_page_break()

    add_heading(doc, "CHƯƠNG 3. DỮ LIỆU VÀ THIẾT KẾ HỆ THỐNG", 1)
    add_heading(doc, "3.1. Bộ dữ liệu sử dụng", 2)
    add_table_df(doc, summarize_inputs(), "Bảng 2. Tổng quan dữ liệu đầu vào của mô hình", "data/*.csv và tham số trong src/")
    add_table_df(doc, macro[["year", "GDP_trillion_VND", "GDP_growth_pct", "digital_economy_share_GDP_pct", "labor_productivity_million_VND"]], "Bảng 3. Một số chỉ tiêu vĩ mô 2020-2025", "reports/figures/00_input_macro.csv")
    add_heading(doc, "3.2. Tiền xử lý dữ liệu", 2)
    add_paragraph(
        doc,
        "Tiền xử lý được thực hiện bằng pandas trong src/data_loader.py và từng module chuyên biệt. Các bước chính "
        "gồm đọc dữ liệu CSV bằng đường dẫn tương đối, kiểm tra tồn tại file, nhận diện linh hoạt tên cột, chuẩn hóa "
        "min-max hoặc vector normalization theo yêu cầu mô hình, và tạo các biến phụ như tỷ trọng phân bổ, điểm chuẩn hóa, "
        "đóng góp tăng trưởng, forecast trajectory hoặc KPI tổng hợp.",
    )
    add_heading(doc, "3.3. Thiết kế pipeline tính toán", 2)
    add_paragraph(
        doc,
        "Pipeline của dự án có dạng: dữ liệu đầu vào -> module xử lý trong src/ -> bảng/biểu đồ trung gian -> dashboard "
        "Streamlit -> khuyến nghị chính sách. Kết quả chạy được xuất ra reports/figures để phục vụ báo cáo, trong đó có "
        "CSV, Markdown và PNG.",
    )
    add_figure(doc, SCREEN_DIR / "01_home.png", "Hình 2. Giao diện tổng quan Bài 12 trên dashboard", 6.3)
    add_heading(doc, "3.4. Thiết kế webapp/dashboard", 2)
    add_paragraph(
        doc,
        "Ứng dụng chính là app.py; mỗi bài có một page riêng trong thư mục pages/. Logic tính toán nằm trong src/ "
        "và được tách khỏi giao diện. Dashboard dùng Streamlit cho webapp, Plotly cho biểu đồ tương tác, pandas/numpy cho "
        "xử lý dữ liệu, SciPy/PuLP/CVXPY hoặc fallback tùy bài toán. Các tác vụ nặng được bọc cache trong page Streamlit.",
    )
    add_table_df(doc, module_table(), "Bảng 4. Tóm tắt 12 module và source code", "Tổng hợp từ thư mục src/ và pages/", max_rows=12)
    add_heading(doc, "3.5. Các tab chức năng của webapp", 2)
    add_paragraph(
        doc,
        "Bài 12 có sáu tab: Tổng quan, 5 kịch bản, Phân bổ, Lao động & AI, Rủi ro và Khuyến nghị. Các screenshot "
        "trong phụ lục E được chụp trực tiếp từ Streamlit local và lưu ở reports/figures/screenshots.",
    )
    add_figure(doc, SCREEN_DIR / "03_budget_allocation.png", "Hình 3. Giao diện kết quả phân bổ ngân sách ở Bài 2", 6.3)
    add_figure(doc, SCREEN_DIR / "04_scenario_comparison.png", "Hình 4. Giao diện so sánh năm kịch bản chính sách", 6.3)
    doc.add_page_break()

    add_heading(doc, "CHƯƠNG 4. KẾT QUẢ CHẠY MÔ HÌNH VÀ WEBAPP", 1)
    add_heading(doc, "4.1. Kết quả module M1: Dự báo kinh tế", 2)
    add_paragraph(
        doc,
        f"Kết quả Bài 1 có MAPE = {mape:.4f} theo reports/figures/00_model_run_summary.csv. "
        f"TFP_A trong bảng fit tăng từ {bai01_fit['TFP_A'].iloc[0]:.4f} năm 2020 lên "
        f"{bai01_fit['TFP_A'].iloc[-1]:.4f} năm 2025. Kịch bản forecast 2030 trong CSV cho "
        f"Y_forecast năm 2030 bằng {bai01_forecast['Y_forecast'].iloc[-1]:,.1f}.",
    )
    add_table_df(doc, bai01_forecast[["year", "K", "L", "D", "AI", "H", "TFP_A", "Y_forecast"]], "Bảng 5. Kết quả dự báo GDP, TFP và biến đầu vào đến năm 2030", "reports/figures/01_bai01_forecast_2030.csv")
    add_figure(doc, FIG_DIR / "01_bai01_gdp_fit.png", "Hình 5. GDP thực tế và GDP dự báo trong Bài 1", 6.0)
    add_figure(doc, FIG_DIR / "01_bai01_forecast_2030.png", "Hình 6. Quỹ đạo GDP forecast đến năm 2030", 6.0)

    add_heading(doc, "4.2. Kết quả module M2: Đánh giá sẵn sàng số", 2)
    add_paragraph(
        doc,
        f"Bài 3 xếp hạng ưu tiên 10 ngành, trong đó ngành đứng đầu theo priority_score là "
        f"{bai03_rank.iloc[0]['sector_name']} với điểm {bai03_rank.iloc[0]['priority_score']:.4f}. "
        f"Bài 6 TOPSIS xếp hạng sáu vùng, vùng đứng đầu là {bai06_topsis.iloc[0]['region_name']} "
        f"với topsis_score {bai06_topsis.iloc[0]['topsis_score']:.4f}.",
    )
    add_table_df(doc, bai06_topsis[["region_name", "topsis_score", "digital_index", "ai_readiness", "trained_labor", "gini", "rank"]], "Bảng 6. Xếp hạng TOPSIS của 6 vùng kinh tế", "reports/figures/06_bai06_topsis_expert.csv")
    add_figure(doc, FIG_DIR / "06_bai06_topsis_score.png", "Hình 7. Biểu đồ điểm TOPSIS theo vùng", 6.0)
    add_table_df(doc, bai03_rank[["rank", "sector_name", "priority_score", "growth", "ai_readiness", "automation_risk"]], "Bảng 7. Top ngành theo priority score", "reports/figures/03_bai03_priority_ranking.csv", max_rows=10)

    add_heading(doc, "4.3. Kết quả module M3: Tối ưu phân bổ ngân sách", 2)
    fair_cost = float(bai04_fair.loc[bai04_fair["scenario"] == "Không fairness", "objective"].iloc[0]) - float(
        bai04_fair.loc[bai04_fair["scenario"] == "Có fairness", "objective"].iloc[0]
    )
    add_paragraph(
        doc,
        "Bài 2 phân bổ ngân sách mặc định 100 nghìn tỷ VND cho bốn hạng mục. Bảng output cho thấy nghiệm tối ưu "
        f"chọn x1={bai02_alloc.loc[bai02_alloc['variable']=='x1','allocation'].iloc[0]:.1f}, "
        f"x2={bai02_alloc.loc[bai02_alloc['variable']=='x2','allocation'].iloc[0]:.1f}, "
        f"x3={bai02_alloc.loc[bai02_alloc['variable']=='x3','allocation'].iloc[0]:.1f}, "
        f"x4={bai02_alloc.loc[bai02_alloc['variable']=='x4','allocation'].iloc[0]:.1f}. "
        f"Trong Bài 4, objective không fairness là {bai04_fair['objective'].iloc[0]:,.2f}, có fairness là "
        f"{bai04_fair['objective'].iloc[1]:,.2f}, do đó chi phí công bằng theo cách tính output là {fair_cost:,.2f}.",
    )
    add_table_df(doc, bai02_alloc, "Bảng 8. Phân bổ tối ưu ngân sách số trong Bài 2", "reports/figures/02_bai02_allocation.csv")
    add_table_df(doc, bai04_matrix, "Bảng 9. Ma trận phân bổ ngân sách tối ưu theo vùng và hạng mục", "reports/figures/04_bai04_allocation_matrix.csv")
    add_figure(doc, FIG_DIR / "04_bai04_allocation_heatmap.png", "Hình 8. Heatmap phân bổ ngân sách theo vùng và hạng mục", 6.0)

    add_heading(doc, "4.4. Kết quả module M4: Mô phỏng lao động", 2)
    top_job = bai09_jobs.sort_values("NetJob", ascending=False).iloc[0]
    add_paragraph(
        doc,
        f"Bài 9 tối ưu phân bổ AI và đào tạo lại theo 8 ngành. Output hiện tại cho thấy ngành có NetJob cao nhất là "
        f"{top_job['sector']} với NetJob {top_job['NetJob']:,.2f}; trong nghiệm này, x_AI = {top_job['x_AI']:,.2f} "
        f"và x_H = {top_job['x_H']:,.2f}. Kết quả này phản ánh cấu trúc hệ số và ràng buộc trong mô hình, không nên "
        "diễn giải như dự báo thị trường lao động thực tế nếu chưa hiệu chỉnh bằng dữ liệu quan sát chi tiết hơn.",
    )
    add_table_df(doc, bai09_jobs[["sector", "labor_million", "risk_pct", "x_AI", "x_H", "NewJob", "DisplacedJob", "RetrainingCapacity", "NetJob"]], "Bảng 10. Kết quả việc làm ròng theo ngành trong Bài 9", "reports/figures/09_bai09_allocation_jobs.csv", max_rows=8)
    add_figure(doc, FIG_DIR / "09_bai09_netjob.png", "Hình 9. Biểu đồ NetJob theo ngành", 6.0)

    add_heading(doc, "4.5. Kết quả module M5: Đánh giá rủi ro", 2)
    add_paragraph(
        doc,
        "Bài 12 cho phép xem đồng thời Inequality_risk, Cyber_risk và Emission_risk theo từng kịch bản. Theo "
        "bảng KPI, S4 Bao trùm số có Inequality_risk thấp nhất trong 5 kịch bản và Cyber_risk thấp nhất; S1 "
        "Truyền thống có Emission_risk cao nhất trong bảng KPI. Các nhận xét này được rút trực tiếp từ "
        "reports/figures/12_bai12_aideom_kpi.csv.",
    )
    add_table_df(doc, bai12_kpi[["scenario", "scenario_name", "Inequality_risk", "Cyber_risk", "Emission_risk", "Overall_score"]], "Bảng 11. Rủi ro và điểm tổng hợp theo kịch bản", "reports/figures/12_bai12_aideom_kpi.csv")
    add_figure(doc, FIG_DIR / "12_bai12_risk_chart.png", "Hình 10. Biểu đồ cảnh báo rủi ro theo kịch bản", 6.0)

    add_heading(doc, "4.6. Kết quả module M6: Webapp/dashboard", 2)
    add_paragraph(
        doc,
        "Webapp cho phép người dùng chọn page, điều chỉnh tham số, chạy mô hình, xem bảng kết quả, xem biểu đồ "
        "Plotly và tải CSV. Bài 12 đóng vai trò trang tổng hợp, tự động chọn kịch bản tốt nhất theo Overall_score "
        "và hiển thị khuyến nghị chính sách.",
    )
    add_figure(doc, SCREEN_DIR / "02_macro_forecast.png", "Hình 11. Giao diện kết quả dự báo/Macro sau khi chạy Bài 1", 6.3)
    add_figure(doc, SCREEN_DIR / "05_risk_warning.png", "Hình 12. Giao diện tab rủi ro trên dashboard tích hợp", 6.3)
    doc.add_page_break()

    add_heading(doc, "CHƯƠNG 5. SO SÁNH 5 KỊCH BẢN CHÍNH SÁCH", 1)
    add_heading(doc, "5.1. Mô tả các kịch bản", 2)
    scen_table = bai12_kpi[["scenario", "scenario_name", "description", "K_share", "D_share", "AI_share", "H_share"]]
    add_table_df(doc, scen_table, "Bảng 12. Mô tả và tỷ trọng phân bổ của 5 kịch bản", "reports/figures/12_bai12_aideom_kpi.csv")
    add_heading(doc, "5.2. So sánh kết quả định lượng", 2)
    add_paragraph(
        doc,
        f"Trong bảng KPI, kịch bản có Overall_score cao nhất là {best_scenario['scenario']} - "
        f"{best_scenario['scenario_name']} với {best_scenario['Overall_score']:.2f}/100. "
        f"Kịch bản có GDP_gain cao nhất là {highest_gdp['scenario']} - {highest_gdp['scenario_name']} "
        f"với {highest_gdp['GDP_gain']:,.2f}. Kịch bản có NetJob cao nhất là "
        f"{highest_netjob['scenario']} - {highest_netjob['scenario_name']} với {highest_netjob['NetJob']:,.2f}.",
    )
    add_table_df(
        doc,
        bai12_kpi[
            [
                "scenario",
                "scenario_name",
                "GDP_gain",
                "Digital_score",
                "AI_score",
                "Human_capital_score",
                "NetJob",
                "Inequality_risk",
                "Cyber_risk",
                "Emission_risk",
                "Overall_score",
            ]
        ],
        "Bảng 13. So sánh KPI của 5 kịch bản chính sách",
        "reports/figures/12_bai12_aideom_kpi.csv",
    )
    add_figure(doc, FIG_DIR / "12_bai12_allocation_heatmap.png", "Hình 13. Heatmap phân bổ K/D/AI/H theo kịch bản", 6.0)
    add_figure(doc, FIG_DIR / "12_bai12_overall_score.png", "Hình 14. Overall_score của 5 kịch bản", 6.0)
    add_heading(doc, "5.3. Phân tích đánh đổi chính sách", 2)
    add_paragraph(
        doc,
        "Các kết quả cho thấy không có kịch bản nào tối ưu tuyệt đối trên mọi tiêu chí. S3 AI dẫn dắt đạt GDP_gain "
        "cao nhất trong bảng KPI nhưng Cyber_risk cũng cao nhất. S4 Bao trùm số đạt NetJob cao nhất và rủi ro thấp hơn "
        "nhưng GDP_gain thấp hơn S3 và S5. S1 Truyền thống có Overall_score thấp nhất trong output tích hợp, chủ yếu do "
        "các chỉ số số hóa, AI, vốn nhân lực và NetJob chuẩn hóa thấp. S5 Tối ưu cân bằng không đứng đầu từng tiêu chí "
        "riêng lẻ nhưng có Overall_score cao nhất nhờ cân bằng nhiều chiều.",
    )
    add_heading(doc, "5.4. Lựa chọn kịch bản khuyến nghị", 2)
    add_paragraph(
        doc,
        "Theo cơ chế điểm tổng hợp trong src/scenario_engine.py và output 12_bai12_aideom_kpi.csv, nhóm khuyến nghị "
        f"xem {best_scenario['scenario']} - {best_scenario['scenario_name']} là phương án nền cho thảo luận chính sách. "
        "Khuyến nghị này cần đi kèm điều kiện triển khai: duy trì đầu tư vào vốn nhân lực, giám sát rủi ro an ninh dữ liệu, "
        "và bảo đảm sàn đầu tư cho các vùng có năng lực hấp thụ thấp hơn.",
    )
    add_table_df(doc, bai12_recs, "Bảng 14. Khuyến nghị chính sách tự động từ dashboard", "reports/figures/12_bai12_policy_recommendations.csv")
    doc.add_page_break()

    add_heading(doc, "CHƯƠNG 6. THẢO LUẬN CHÍNH SÁCH", 1)
    add_heading(doc, "6.1. Ý nghĩa kết quả đối với Việt Nam", 2)
    add_paragraph(
        doc,
        "Kết quả của AIDEOM-VN gợi ý rằng chính sách chuyển đổi số cần phối hợp giữa đầu tư hạ tầng, dữ liệu, AI và "
        "vốn nhân lực. Kịch bản chỉ ưu tiên vốn truyền thống không tạo điểm tổng hợp cao trong output Bài 12. Ngược lại, "
        "các kịch bản có thành phần số hóa, AI và nhân lực lớn hơn cải thiện các KPI liên quan đến Digital_score, AI_score "
        "và NetJob, nhưng cũng làm xuất hiện đánh đổi về cyber risk hoặc chi phí công bằng.",
    )
    add_heading(doc, "6.2. Vai trò của AI trong hỗ trợ ra quyết định", 2)
    add_paragraph(
        doc,
        "AI và mô hình định lượng trong dự án đóng vai trò hỗ trợ phân tích, không thay thế quyết định chính trị - xã hội. "
        "Dashboard giúp người dùng quan sát mối quan hệ giữa giả định và kết quả, phát hiện ràng buộc binding, so sánh kịch bản "
        "và tải dữ liệu kết quả. Tuy nhiên, việc quyết định chính sách thật cần thêm tham vấn chuyên gia, dữ liệu hành chính, "
        "nguồn lực ngân sách thực tế và đánh giá tác động phân phối.",
    )
    add_heading(doc, "6.3. Hạn chế của mô hình", 2)
    add_bullets(
        doc,
        [
            "Dữ liệu đầu vào có quy mô nhỏ: 6 năm vĩ mô, 10 ngành và 6 vùng.",
            "Một số tham số mô hình là giả định được hard-code trong source code để phục vụ nguyên mẫu.",
            "Mô hình chưa phản ánh đầy đủ hành vi doanh nghiệp, hộ gia đình và phản ứng chính sách động.",
            "Các output là kết quả mô phỏng theo đơn vị và hàm mục tiêu của từng bài, không phải dự báo chính thức.",
            "Webapp mới là nguyên mẫu học thuật, chưa có kết nối dữ liệu thời gian thực hoặc phân quyền người dùng.",
        ],
    )
    add_heading(doc, "6.4. Đề xuất cải tiến", 2)
    add_bullets(
        doc,
        [
            "Cập nhật dữ liệu theo quý và bổ sung dữ liệu địa phương/tiểu ngành nếu có nguồn chính thức.",
            "Hiệu chỉnh hệ số bằng dữ liệu quan sát, thay vì dùng toàn bộ hệ số giả định.",
            "Bổ sung bản đồ địa lý tương tác và bộ lọc theo vùng/ngành.",
            "Mở rộng mô hình học tăng cường và kiểm định chính sách dưới nhiều shock vĩ mô.",
            "Kết nối dashboard với quy trình xuất báo cáo tự động để lưu vết giả định, kết quả và phiên bản mô hình.",
        ],
    )
    add_heading(doc, "6.5. Đối chiếu một số mô hình bổ sung", 2)
    add_table_df(doc, bai07_compare, "Bảng 15. So sánh nghiệm tăng trưởng cao nhất và nghiệm thỏa hiệp Pareto", "reports/figures/07_bai07_growth_vs_compromise.csv")
    add_table_df(doc, bai08_summary, "Bảng 16. So sánh chiến lược trong mô hình tối ưu động", "reports/figures/08_bai08_strategy_summary.csv")
    add_table_df(doc, bai10_vss, "Bảng 17. VSS và EVPI trong quy hoạch ngẫu nhiên", "reports/figures/10_bai10_vss_evpi.csv")
    add_table_df(doc, bai11_compare, "Bảng 18. So sánh reward trong Q-learning", "reports/figures/11_bai11_policy_comparison.csv")
    doc.add_page_break()

    add_heading(doc, "CHƯƠNG 7. KẾT LUẬN VÀ KHUYẾN NGHỊ", 1)
    add_heading(doc, "7.1. Kết luận chính", 2)
    add_paragraph(
        doc,
        "Đồ án đã xây dựng một dashboard Streamlit gồm 12 module ra quyết định, tách logic tính toán trong src/ và "
        "giao diện trong pages/. Hệ thống đọc ba bộ dữ liệu đầu vào, chạy các mô hình dự báo, xếp hạng, tối ưu, mô phỏng "
        "lao động, quy hoạch ngẫu nhiên, học tăng cường và tổng hợp năm kịch bản chính sách.",
    )
    add_paragraph(
        doc,
        f"Kết quả tích hợp cho thấy {best_scenario['scenario']} - {best_scenario['scenario_name']} có Overall_score cao nhất "
        f"({best_scenario['Overall_score']:.2f}/100), trong khi {highest_gdp['scenario']} - {highest_gdp['scenario_name']} "
        "có GDP_gain cao nhất và S4 Bao trùm số nổi bật về NetJob cũng như rủi ro thấp hơn. Đây là cơ sở để thảo luận "
        "đánh đổi chính sách thay vì chọn phương án theo một mục tiêu duy nhất.",
    )
    add_heading(doc, "7.2. Khuyến nghị chính sách", 2)
    add_bullets(
        doc,
        [
            "Ưu tiên đầu tư nhân lực số trước hoặc đồng thời với mở rộng AI quy mô lớn.",
            "Tập trung AI ở vùng có năng lực hấp thụ cao nhưng duy trì sàn đầu tư cho vùng yếu hơn.",
            "Kết hợp tăng trưởng với bao trùm số để tránh rủi ro bất bình đẳng vùng.",
            "Thiết lập cơ chế giám sát rủi ro dữ liệu, lao động và môi trường trên dashboard.",
        ],
    )
    add_heading(doc, "7.3. Hướng phát triển tiếp theo", 2)
    add_paragraph(
        doc,
        "AIDEOM-VN có thể phát triển thành đề tài nghiên cứu sâu hơn, dashboard thử nghiệm cho lớp học hoặc công cụ hỗ trợ "
        "hoạch định chính sách ở cấp mô phỏng. Để tiến tới ứng dụng thực tế, cần bổ sung dữ liệu chính thức, kiểm định mô hình, "
        "mở rộng giao diện bản đồ và xây dựng quy trình cập nhật kết quả định kỳ.",
    )
    doc.add_page_break()

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_paragraph(
        doc,
        "Các tài liệu tham khảo dưới đây được liệt kê như nền tảng phương pháp và nguồn dữ liệu cần đối chiếu khi phát triển "
        "phiên bản nghiên cứu hoàn chỉnh. Báo cáo hiện tại không sử dụng số liệu ngoài các tệp local của dự án.",
    )
    add_bullets(
        doc,
        [
            "Tổng cục Thống kê/NSO-GSO: dữ liệu kinh tế - xã hội Việt Nam.",
            "World Bank: dữ liệu kinh tế vĩ mô và phát triển số.",
            "Bộ Thông tin và Truyền thông: chuyển đổi số, kinh tế số, chỉ số số.",
            "Bộ Kế hoạch và Đầu tư: thông tin đầu tư, FDI và phát triển vùng.",
            "Hwang, C. L. và Yoon, K.: Multiple Attribute Decision Making, nền tảng phương pháp TOPSIS.",
            "Sutton, R. S. và Barto, A. G.: Reinforcement Learning, nền tảng Q-learning.",
            "Tài liệu về linear programming, mixed-integer programming, stochastic programming và multi-objective optimization.",
        ],
    )
    doc.add_page_break()

    add_heading(doc, "PHỤ LỤC A. MÃ NGUỒN CHÍNH", 1)
    add_paragraph(
        doc,
        "Mã nguồn chính nằm trong src/. Các module được tách khỏi giao diện để bảo đảm khả năng kiểm thử và tái sử dụng. "
        "Một số file quan trọng gồm src/data_loader.py, src/bai01_cobb_douglas.py, src/bai02_lp_budget.py, "
        "src/bai04_region_lp.py, src/bai09_labor_ai.py và src/scenario_engine.py.",
    )
    add_table_df(doc, run_summary, "Bảng 19. Tóm tắt kết quả chạy 12 bài", "reports/figures/00_model_run_summary.csv")

    add_heading(doc, "PHỤ LỤC B. CẤU TRÚC THƯ MỤC DỰ ÁN", 1)
    add_paragraph(
        doc,
        "Cấu trúc chính của dự án gồm app.py, data/, pages/, src/, outputs/, tests/, reports/, requirements.txt và README.md. "
        "Các bảng/hình dùng trong báo cáo nằm dưới reports/figures.",
    )
    add_paragraph(
        doc,
        "aideom_vn/\n├── app.py\n├── data/\n├── pages/\n├── src/\n├── outputs/\n├── reports/\n├── tests/\n├── requirements.txt\n└── README.md",
    )

    add_heading(doc, "PHỤ LỤC C. HƯỚNG DẪN CHẠY WEBAPP", 1)
    add_paragraph(doc, "Cài đặt và chạy local trên macOS/Python 3.10+:", bold=True)
    add_paragraph(doc, "python -m venv .venv\nsource .venv/bin/activate\npip install -r requirements.txt\nstreamlit run app.py")
    add_paragraph(
        doc,
        "Sau khi chạy, mở trình duyệt tại http://localhost:8501. Người dùng có thể chọn từng bài ở sidebar, điều chỉnh "
        "tham số, chạy mô hình, xem biểu đồ và tải CSV kết quả.",
    )

    add_heading(doc, "PHỤ LỤC D. BẢNG KẾT QUẢ ĐẦY ĐỦ", 1)
    add_table_df(doc, bai05_selected[["project_id", "project_name", "sector", "cost", "benefit", "selected_label"]], "Bảng 20. Các dự án được chọn trong Bài 5", "reports/figures/05_bai05_selected_projects.csv")
    add_table_df(doc, bai12_alloc, "Bảng 21. Phân bổ chi tiết của 5 kịch bản theo K/D/AI/H", "reports/figures/12_bai12_allocation_long.csv", max_rows=20)

    add_heading(doc, "PHỤ LỤC E. ẢNH CHỤP GIAO DIỆN WEBAPP", 1)
    add_figure(doc, SCREEN_DIR / "01_home.png", "Hình E1. Trang tổng quan Bài 12", 6.3)
    add_figure(doc, SCREEN_DIR / "02_macro_forecast.png", "Hình E2. Trang dữ liệu/dự báo kinh tế", 6.3)
    add_figure(doc, SCREEN_DIR / "03_budget_allocation.png", "Hình E3. Trang chạy mô hình phân bổ ngân sách", 6.3)
    add_figure(doc, SCREEN_DIR / "04_scenario_comparison.png", "Hình E4. Trang so sánh kịch bản", 6.3)
    add_figure(doc, SCREEN_DIR / "05_risk_warning.png", "Hình E5. Trang cảnh báo rủi ro", 6.3)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_report()
